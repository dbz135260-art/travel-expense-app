"""Regression tests for certificate generation bug fixes.

Covers the following bugs fixed in recent commits:

1. Date format mismatch — generate_certificate_word built a date range
   like "2026年04月15日-2026年04月17日" (year repeated), but the template
   uses "2026年04月15日-04月17日" (year only on the first date).  The old
   code's string never matched the template, so replacement silently
   failed.

2. Sign-off paragraph mis-identification — the old search ("年" in text
   and "月" in text and "中国" not in text) could match the long
   training-text paragraph instead of the short sign-off date line.

3. Sign-off alignment — the sign-off date paragraph must be RIGHT-
   aligned.

4. _set_cell_text — the old code directly accessed
   cell.paragraphs[0].runs[0].text which crashed on cells with no runs
   and left stale text in extra paragraphs.  The new helper clears all
   paragraphs before writing.
"""

import io
import sys
from datetime import date
from pathlib import Path

import pytest

# Ensure the repo root is importable
sys.path.insert(0, str(Path(__file__).resolve().parent))

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app import _set_cell_text, generate_approval_word, generate_certificate_word


# ────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────

TEMPLATE_DIR = Path(__file__).resolve().parent


def _load_doc(raw_bytes: bytes) -> docx.Document:
    return docx.Document(io.BytesIO(raw_bytes))


def _make_params(
    start: date = date(2026, 4, 15),
    end: date = date(2026, 4, 17),
    project: str = "测试项目",
    hours: str = "32",
    location: str = "北京",
    person: str = "张三",
    cert_year: str = "2026",
    cert_seq: str = "0001",
) -> dict:
    return {
        "project_name": project,
        "train_location": location,
        "hours": hours,
        "cert_year": cert_year,
        "cert_seq_start": cert_seq,
        "person_in_charge": person,
        "train_date_start": start,
        "train_date_end": end,
        "train_date": f"{start.year}年{start.month}月{start.day}日-{end.month}月{end.day}日",
        "cert_date_str": f"{cert_year}年04月17日",
    }


# ────────────────────────────────────────────────────────────────────
# Bug 1: Date format must match the template pattern
# ────────────────────────────────────────────────────────────────────


class TestCertificateDateFormatting:
    """The template contains dates like '2026年04月15日-04月17日'.

    The old code built '2026年04月15日-2026年04月17日' (year on both sides)
    and tried to replace it, which never matched.  The fix constructs
    old_date with the year only on the first date so the replacement
    succeeds.
    """

    def test_date_replaced_with_unpadded_format(self):
        """After generation, the training paragraph should contain the
        *new* un-padded date (e.g. '4月15日') instead of the original
        zero-padded template date ('04月15日')."""
        params = _make_params(start=date(2026, 4, 15), end=date(2026, 4, 17))
        doc = _load_doc(generate_certificate_word(params))

        training_para = None
        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            if "参加" in full or "姓名" in full:
                training_para = full
                break

        assert training_para is not None, "Training text paragraph not found"
        # The old zero-padded template date should be gone
        assert "04月15日" not in training_para, (
            "Zero-padded template date was NOT replaced — "
            "date format mismatch bug has regressed"
        )
        # The new un-padded date should be present
        assert "4月15日" in training_para
        assert "4月17日" in training_para

    def test_date_replacement_does_not_duplicate_year(self):
        """The replacement must NOT produce '2026年…-2026年…' (year on
        both halves).  This was the root cause of the old bug: the
        replace target included the year twice, so it never matched the
        template."""
        params = _make_params(start=date(2026, 4, 15), end=date(2026, 4, 17))
        doc = _load_doc(generate_certificate_word(params))

        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            if "参加" in full:
                # Must not have a doubled year
                assert full.count("2026年") == 1, (
                    f"Year appears more than once in the date range: {full!r}"
                )
                break

    def test_project_name_replaced(self):
        """Verify the project name is injected into the training text."""
        params = _make_params(project="回归测试项目名称")
        doc = _load_doc(generate_certificate_word(params))

        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            if "参加" in full or "回归测试" in full:
                assert "回归测试项目名称" in full
                break
        else:
            pytest.fail("Project name not found in any paragraph")

    def test_hours_replaced(self):
        """Verify training hours are replaced."""
        params = _make_params(hours="48")
        doc = _load_doc(generate_certificate_word(params))

        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            if "学时" in full:
                assert "48学时" in full
                # Original template value should be gone
                assert "20学时" not in full
                break


# ────────────────────────────────────────────────────────────────────
# Bug 2 & 3: Sign-off paragraph identification and alignment
# ────────────────────────────────────────────────────────────────────


class TestSignOffParagraph:
    """The sign-off date paragraph is a short line like '2026年04月'.

    Bug 2: The old search matched any paragraph containing '年' and '月'
    (excluding '中国'), which could hit the long training paragraph.
    Bug 3: The sign-off paragraph must be RIGHT-aligned.
    """

    def test_sign_off_is_right_aligned(self):
        """The sign-off date paragraph must have RIGHT alignment."""
        params = _make_params()
        doc = _load_doc(generate_certificate_word(params))

        sign_off_para = None
        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            txt = full.strip()
            if txt and "年" in txt and "月" in txt and len(txt) <= 10 and "中国" not in txt:
                sign_off_para = p
                break

        assert sign_off_para is not None, "Sign-off paragraph not found"
        assert sign_off_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.RIGHT, (
            f"Sign-off paragraph alignment is "
            f"{sign_off_para.paragraph_format.alignment}, expected RIGHT"
        )

    def test_sign_off_not_same_as_training_paragraph(self):
        """The sign-off paragraph must be distinct from the training text
        paragraph — the old code could accidentally overwrite the training
        paragraph if it was found first."""
        params = _make_params()
        doc = _load_doc(generate_certificate_word(params))

        training_para = None
        sign_off_para = None
        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            if "参加" in full or "姓名" in full:
                training_para = p
            txt = full.strip()
            if txt and "年" in txt and "月" in txt and len(txt) <= 10 and "中国" not in txt:
                sign_off_para = p

        assert training_para is not None
        assert sign_off_para is not None
        assert training_para is not sign_off_para, (
            "Sign-off paragraph should be distinct from training paragraph"
        )

    def test_sign_off_contains_correct_year_month(self):
        """The sign-off should reflect the training end month."""
        params = _make_params(end=date(2026, 5, 20))
        doc = _load_doc(generate_certificate_word(params))

        for p in doc.paragraphs:
            full = "".join(r.text for r in p.runs)
            txt = full.strip()
            if txt and "年" in txt and "月" in txt and len(txt) <= 10 and "中国" not in txt:
                assert "2026年05月" in txt, f"Expected '2026年05月' in sign-off, got {txt!r}"
                break
        else:
            pytest.fail("Sign-off paragraph not found")


# ────────────────────────────────────────────────────────────────────
# Bug 4: _set_cell_text clears all paragraphs
# ────────────────────────────────────────────────────────────────────


class TestSetCellText:
    """_set_cell_text must clear ALL paragraphs in a cell, not just the
    first run.  The old code accessed cell.paragraphs[0].runs[0].text
    directly, which:
      - crashed on cells with no runs
      - left stale text in extra paragraphs
    """

    def _make_cell_with_runs(self, *texts):
        """Create a minimal docx with a table cell containing given texts
        in separate runs of the first paragraph."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        # Clear default paragraph
        cell.paragraphs[0].clear()
        for t in texts:
            cell.paragraphs[0].add_run(t)
        return cell

    def _make_cell_with_extra_paragraphs(self):
        """Create a cell with multiple paragraphs containing text."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].clear()
        cell.paragraphs[0].add_run("first paragraph text")
        # python-docx: adding a paragraph to a cell appends a new <w:p>
        cell.add_paragraph("second paragraph text")
        return cell

    def _make_cell_no_runs(self):
        """Create a cell whose first paragraph has no runs at all."""
        doc = docx.Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.paragraphs[0].clear()
        return cell

    def test_single_run_replaced(self):
        cell = self._make_cell_with_runs("old text")
        _set_cell_text(cell, "new text")
        result = "".join(r.text for p in cell.paragraphs for r in p.runs)
        assert result == "new text"

    def test_multiple_runs_cleared(self):
        cell = self._make_cell_with_runs("aaa", "bbb", "ccc")
        _set_cell_text(cell, "replaced")
        result = "".join(r.text for p in cell.paragraphs for r in p.runs)
        assert result == "replaced"

    def test_extra_paragraphs_cleared(self):
        """Stale text in extra paragraphs must be cleared."""
        cell = self._make_cell_with_extra_paragraphs()
        _set_cell_text(cell, "only this")
        result = "".join(r.text for p in cell.paragraphs for r in p.runs)
        assert result == "only this", (
            f"Extra paragraph text was not cleared: {result!r}"
        )

    def test_no_runs_adds_run(self):
        """If the first paragraph has no runs, _set_cell_text must add one."""
        cell = self._make_cell_no_runs()
        _set_cell_text(cell, "brand new")
        result = "".join(r.text for p in cell.paragraphs for r in p.runs)
        assert result == "brand new"


# ────────────────────────────────────────────────────────────────────
# Integration: generate_approval_word uses _set_cell_text
# ────────────────────────────────────────────────────────────────────


class TestApprovalWord:
    """generate_approval_word fills a table via _set_cell_text.

    The old code directly set runs, which could crash or leave stale text.
    These tests verify the end-to-end output is correct.
    """

    def test_project_name_filled(self):
        params = _make_params(project="审批测试项目")
        raw = generate_approval_word(params, student_count=10)
        doc = _load_doc(raw)
        cell_text = "".join(
            r.text for p in doc.tables[0].rows[0].cells[1].paragraphs for r in p.runs
        )
        assert "审批测试项目" in cell_text

    def test_date_range_filled(self):
        params = _make_params(start=date(2026, 4, 15), end=date(2026, 4, 17))
        raw = generate_approval_word(params, student_count=5)
        doc = _load_doc(raw)
        cell_text = "".join(
            r.text for p in doc.tables[0].rows[1].cells[1].paragraphs for r in p.runs
        )
        assert "2026.04.15-2026.04.17" in cell_text

    def test_student_count_filled(self):
        params = _make_params()
        raw = generate_approval_word(params, student_count=42)
        doc = _load_doc(raw)
        cell_text = "".join(
            r.text for p in doc.tables[0].rows[2].cells[4].paragraphs for r in p.runs
        )
        assert "42" in cell_text

    def test_cert_range_filled(self):
        params = _make_params(cert_year="2026", cert_seq="0100")
        raw = generate_approval_word(params, student_count=10)
        doc = _load_doc(raw)
        cell_text = "".join(
            r.text for p in doc.tables[0].rows[6].cells[0].paragraphs for r in p.runs
        )
        assert "CASEI-HYPX-2026-0100" in cell_text
        assert "CASEI-HYPX-2026-0109" in cell_text

    def test_no_crash_on_template_cells(self):
        """Simply calling generate_approval_word should not raise,
        even when template cells have merged cells or multiple paragraphs."""
        params = _make_params()
        raw = generate_approval_word(params, student_count=1)
        assert len(raw) > 0
