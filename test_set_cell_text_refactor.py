"""
Verify _set_cell_text refactor (commit 30c6274) produces correct output.

The refactor changed _set_cell_text from clearing only the first paragraph's
runs to clearing ALL paragraphs' runs in a cell. This was a bug fix: the old
version left stale template text in multi-paragraph cells (e.g. Row 6 of the
approval template, which has 2 paragraphs for the cert range).

This test generates approval documents with multiple data sets and verifies:
  1. Single-paragraph cells: identical output between old and new.
  2. Multi-paragraph cells: new version correctly clears stale text.
"""
import io
import docx
from pathlib import Path
from datetime import date

TEMPLATE_APPROVAL = Path(__file__).parent / "审批表模板.docx"


def _set_cell_text_old(cell, text):
    """Pre-refactor: only clears first paragraph runs."""
    p = cell.paragraphs[0]
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def _set_cell_text_new(cell, text):
    """Post-refactor: clears ALL paragraphs' runs."""
    for i, p in enumerate(cell.paragraphs):
        for r in p.runs:
            r.text = ""
        if i == 0:
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)


def cell_text(cell):
    """All paragraph text joined by |."""
    return "|".join("".join(r.text for r in p.runs) for p in cell.paragraphs)


def build_approval(set_cell_fn, params, student_count):
    train_start = params["train_date_start"]
    train_end = params["train_date_end"]
    start_str = f"{train_start.year}.{train_start.month:02d}.{train_start.day:02d}"
    end_str = f"{train_end.year}.{train_end.month:02d}.{train_end.day:02d}"
    date_range = f"{start_str}-{end_str}"
    cert_year = params.get("cert_year", "2026")
    cert_start = int(params.get("cert_seq_start", "0001"))
    cert_end = cert_start + student_count - 1
    cert_range = (
        f"CASEI-HYPX-{cert_year}-{str(cert_start).zfill(4)}  --  "
        f"CASEI-HYPX-{cert_year}-{str(cert_end).zfill(4)}"
    )

    doc = docx.Document(str(TEMPLATE_APPROVAL))
    table = doc.tables[0]
    set_cell_fn(table.rows[0].cells[1], params.get("project_name", ""))
    set_cell_fn(table.rows[1].cells[1], date_range)
    set_cell_fn(table.rows[1].cells[4], params.get("train_location", ""))
    set_cell_fn(table.rows[2].cells[1], params.get("person_in_charge", ""))
    set_cell_fn(table.rows[2].cells[4], str(student_count))
    set_cell_fn(table.rows[6].cells[0], f"证书号码（号段）\n{cert_range}")
    return doc, table


SCENARIOS = [
    {
        "name": "standard",
        "params": {
            "project_name": "起重机械安全培训项目",
            "train_location": "北京市朝阳区",
            "cert_year": "2026", "cert_seq_start": "0001",
            "person_in_charge": "张三",
            "train_date_start": date(2026, 4, 15),
            "train_date_end": date(2026, 4, 17),
        },
        "student_count": 45,
    },
    {
        "name": "large_batch",
        "params": {
            "project_name": "第五届全国特种设备检验检测培训",
            "train_location": "上海浦东",
            "cert_year": "2026", "cert_seq_start": "0100",
            "person_in_charge": "李四",
            "train_date_start": date(2026, 6, 1),
            "train_date_end": date(2026, 6, 5),
        },
        "student_count": 200,
    },
    {
        "name": "minimal",
        "params": {
            "project_name": "A",
            "train_location": "B",
            "cert_year": "2026", "cert_seq_start": "9999",
            "person_in_charge": "C",
            "train_date_start": date(2026, 1, 1),
            "train_date_end": date(2026, 1, 1),
        },
        "student_count": 1,
    },
    {
        "name": "empty_strings",
        "params": {
            "project_name": "",
            "train_location": "",
            "cert_year": "2026", "cert_seq_start": "0001",
            "person_in_charge": "",
            "train_date_start": date(2026, 3, 10),
            "train_date_end": date(2026, 3, 12),
        },
        "student_count": 10,
    },
]


def test_single_paragraph_cells_unchanged():
    """Single-paragraph cells produce identical text with old and new."""
    single_para_rows = [
        (0, 1, "project_name"),
        (1, 1, "date_range"),
        (1, 4, "location"),
        (2, 1, "person"),
        (2, 4, "count"),
    ]
    for sc in SCENARIOS:
        _, t_old = build_approval(_set_cell_text_old, sc["params"], sc["student_count"])
        _, t_new = build_approval(_set_cell_text_new, sc["params"], sc["student_count"])
        for row, col, label in single_para_rows:
            old_t = cell_text(t_old.rows[row].cells[col])
            new_t = cell_text(t_new.rows[row].cells[col])
            assert old_t == new_t, (
                f"[{sc['name']}] {label} mismatch: old={old_t!r} new={new_t!r}"
            )


def test_multi_paragraph_cell_stale_text_cleared():
    """Row 6 (cert range) has 2 paragraphs; new version clears stale P1."""
    template_p1 = "CASEI-HYPX-2026-0827  --   CASEI-HYPX-2026-0994"

    for sc in SCENARIOS:
        _, t_old = build_approval(_set_cell_text_old, sc["params"], sc["student_count"])
        _, t_new = build_approval(_set_cell_text_new, sc["params"], sc["student_count"])

        old_cell = t_old.rows[6].cells[0]
        new_cell = t_new.rows[6].cells[0]

        # Old version leaves stale template text in paragraph 1
        old_p1 = "".join(r.text for r in old_cell.paragraphs[1].runs)
        assert old_p1 == template_p1, (
            f"[{sc['name']}] expected old P1 to retain template text, got {old_p1!r}"
        )

        # New version clears paragraph 1
        new_p1 = "".join(r.text for r in new_cell.paragraphs[1].runs)
        assert new_p1 == "", (
            f"[{sc['name']}] expected new P1 to be empty, got {new_p1!r}"
        )

        # Both have the correct text in paragraph 0
        old_p0 = "".join(r.text for r in old_cell.paragraphs[0].runs)
        new_p0 = "".join(r.text for r in new_cell.paragraphs[0].runs)
        assert old_p0 == new_p0, (
            f"[{sc['name']}] P0 text should match: old={old_p0!r} new={new_p0!r}"
        )


def test_new_version_roundtrip_clean():
    """New version produces clean output after save-and-reload."""
    for sc in SCENARIOS:
        doc, _ = build_approval(_set_cell_text_new, sc["params"], sc["student_count"])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        doc2 = docx.Document(buf)
        cell = doc2.tables[0].rows[6].cells[0]
        p1_text = "".join(r.text for r in cell.paragraphs[1].runs)
        assert p1_text == "", (
            f"[{sc['name']}] after roundtrip, P1 should be empty, got {p1_text!r}"
        )


def test_empty_run_cell():
    """Both versions handle cells with no existing runs (add_run path)."""
    doc1 = docx.Document()
    t1 = doc1.add_table(rows=1, cols=1)
    _set_cell_text_old(t1.rows[0].cells[0], "Hello")

    doc2 = docx.Document()
    t2 = doc2.add_table(rows=1, cols=1)
    _set_cell_text_new(t2.rows[0].cells[0], "Hello")

    assert cell_text(t1.rows[0].cells[0]) == cell_text(t2.rows[0].cells[0]) == "Hello"


if __name__ == "__main__":
    test_single_paragraph_cells_unchanged()
    print("PASS: single-paragraph cells unchanged")

    test_multi_paragraph_cell_stale_text_cleared()
    print("PASS: multi-paragraph cell stale text cleared by new version")

    test_new_version_roundtrip_clean()
    print("PASS: new version roundtrip clean")

    test_empty_run_cell()
    print("PASS: empty-run cell handled identically")

    print("\nAll verification tests passed.")
