import streamlit as st
import json
import io
import zipfile
import re
import xml.etree.ElementTree as ET
from datetime import datetime, date
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import fitz
import requests
import base64

st.set_page_config(page_title="差旅/劳务费处理", page_icon="🧾", layout="wide")

if "work_log" not in st.session_state:
    st.session_state.work_log = []

QWEN_API_URL = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
QWEN_MODEL = "qwen3.7-plus"

# ─── Qwen3.7-Plus unified API (text + image) ────────

def call_qwen(api_key: str, system_prompt: str, user_prompt: str,
              image_base64: str = "", max_tokens: int = 4096) -> str:
    """Call Qwen3.7-Plus with optional image input."""
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    content = []
    if image_base64:
        content.append({"type": "image_url", "image_url": f"data:image/png;base64,{image_base64}"})
    content.append({"type": "text", "text": user_prompt})
    payload = {
        "model": QWEN_MODEL,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": content}
        ],
        "temperature": 0.01,
        "max_tokens": max_tokens
    }
    resp = requests.post(QWEN_API_URL, headers=headers, json=payload, timeout=120)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def _pdf_to_image_base64(file_bytes: bytes, page_num: int = 0) -> str:
    """Render a PDF page as a PNG image, return base64."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc[page_num]
    pix = page.get_pixmap(dpi=200)
    img_bytes = pix.tobytes("png")
    doc.close()
    return base64.b64encode(img_bytes).decode("utf-8")


def _has_amount_info(text: str) -> bool:
    """Check if extracted text contains any amount/number near amount keywords."""
    amount_kw = ['合计', '票价', '金额', 'CNY', '¥', '￥', '元', '燃油', '民航']
    for kw in amount_kw:
        idx = text.find(kw)
        if idx >= 0:
            nearby = text[max(0, idx):idx + 150]
            if re.search(r'\d+\.?\d*', nearby):
                return True
    return False


def smart_extract_pdf(file_bytes: bytes, qwen_api_key: str = "") -> str:
    """Extract text, fall back to Qwen-VL OCR if text layer is missing amounts."""
    text = extract_pdf_text(file_bytes)
    if _has_amount_info(text):
        return text
    if not qwen_api_key:
        return text
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = len(doc)
        doc.close()
        texts = []
        for p in range(pages):
            img_b64 = _pdf_to_image_base64(file_bytes, p)
            ocr_text = call_qwen(qwen_api_key,
                "你是一个发票OCR助手，逐字识别图片中的文字。",
                "逐字识别这张图片上的所有文字，保持数字和中文原样输出。",
                image_base64=img_b64, max_tokens=2048)
            texts.append(ocr_text)
        combined = "\n".join(texts)
        if combined.strip():
            return combined
    except Exception:
        pass
    return text

def parse_ofd(file_bytes: bytes) -> Dict:
    import zipfile
    result = {"姓名": "", "金额": 0.0, "日期": "", "保险费": 0.0, "类型": "ofd", "原始文本": ""}
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            xml_files = [n for n in z.namelist() if n.endswith(".xml") and "Attachs" in n]
            doc_xmls = [n for n in z.namelist() if n.endswith("Document.xml")]

            for xml_path in xml_files + doc_xmls:
                with z.open(xml_path) as f:
                    raw = f.read()
                    result["原始文本"] += raw.decode("utf-8", errors="replace") + "\n"

            for xml_path in xml_files:
                with z.open(xml_path) as f:
                    raw = f.read()
                    try:
                        root = ET.fromstring(raw)
                        ns = {}
                        for m in re.finditer(r'xmlns:(\w+)="([^"]+)"', raw.decode("utf-8", errors="replace")):
                            ns[m.group(1)] = m.group(2)

                        # Passenger name
                        for el in root.iter():
                            if "PassengerName" in el.tag and el.text:
                                result["姓名"] = el.text.strip()
                                break

                        # Total amount
                        for el in root.iter():
                            if "TotalAmount" in el.tag and el.text:
                                try: result["金额"] = float(el.text)
                                except: pass
                                break

                        # Insurance / other taxes
                        for el in root.iter():
                            if "OtherTaxes" in el.tag and el.text:
                                try: result["保险费"] = float(el.text)
                                except: pass

                        # Date
                        for el in root.iter():
                            if "CarrierDate" in el.tag and el.text:
                                result["日期"] = el.text.strip()
                                break
                            if "IssuanceDate" in el.tag and el.text:
                                if not result["日期"]:
                                    result["日期"] = el.text.strip()
                    except ET.ParseError:
                        continue
    except Exception as e:
        result["原始文本"] = f"OFD解析错误: {str(e)}"
    return result


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        for page in doc:
            text = page.get_text().strip()
            if text:
                text_parts.append(text)
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        return f"[PDF解析错误: {str(e)}]"


# ─── OCR via Qwen-VL API (zero extra dependencies) ──

QWEN_OCR_PROMPT = """你是一个发票OCR助手。从图片中逐字识别所有文字，保持数字和中文原样输出。

特别关注：
1. 姓名、旅客姓名
2. 金额（票价、合计、实付金额等）
3. 日期
4. 发票号码
5. 任何数字数据

输出格式：直接把识别到的文字按行输出。"""

def _has_amount_info(text: str) -> bool:
    """Check if extracted text contains any amount/number near amount keywords."""
    amount_kw = ['合计', '票价', '金额', 'CNY', '¥', '￥', '元', '燃油', '民航']
    for kw in amount_kw:
        idx = text.find(kw)
        if idx >= 0:
            nearby = text[max(0, idx):idx + 150]
            if re.search(r'\d+\.?\d*', nearby):
                return True
    return False


def _pdf_to_image_base64(file_bytes: bytes, page_num: int = 0) -> str:
    """Render a PDF page as a PNG image, return base64."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    page = doc[page_num]
    pix = page.get_pixmap(dpi=200)
    img_bytes = pix.tobytes("png")
    doc.close()
    return base64.b64encode(img_bytes).decode("utf-8")


def _call_qwen_ocr(image_base64: str, api_key: str) -> str:
    """Call Qwen-VL API to do OCR on a base64 image."""
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {
        "model": QWEN_VL_MODEL,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": f"data:image/png;base64,{image_base64}"},
                {"type": "text", "text": QWEN_OCR_PROMPT}
            ]
        }],
        "temperature": 0.01
    }
    resp = requests.post(QWEN_VL_API_URL, headers=headers, json=payload, timeout=60)
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def smart_extract_pdf(file_bytes: bytes, qwen_api_key: str = "") -> str:
    """Extract text, fall back to Qwen-VL OCR if text layer is missing amounts."""
    # Step 1: try text extraction
    text = extract_pdf_text(file_bytes)
    if _has_amount_info(text):
        return text

    # Step 2: try OCR via Qwen-VL (if API key provided)
    if not qwen_api_key:
        return text
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = len(doc)
        doc.close()
        texts = []
        for p in range(pages):
            img_b64 = _pdf_to_image_base64(file_bytes, p)
            ocr_result = _call_qwen_ocr(img_b64, qwen_api_key)
            texts.append(ocr_result)
        combined = "\n".join(texts)
        if combined.strip():
            return combined
    except Exception:
        pass
    return text


def parse_json_response(content: str) -> List[Dict]:
    text = content.strip()
    if text.startswith("```"):
        lines = text.split("\n")
        lines = [l for l in lines if not l.startswith("```")]
        text = "\n".join(lines).strip()
    return json.loads(text)


def create_renamed_zip(file_map: Dict[str, bytes], rename_map: Dict[str, str]) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for orig_name, file_bytes in file_map.items():
            new_name = rename_map.get(orig_name, orig_name)
            zf.writestr(new_name, file_bytes)
    buf.seek(0)
    return buf.getvalue()


def log_work(entry_type, summary, details=None):
    """Add an entry to the work log"""
    from datetime import datetime
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    entry = {"time": now, "type": entry_type, "summary": summary, "details": details or {}}
    if "work_log" in st.session_state:
        st.session_state.work_log.append(entry)


def merge_pdfs(file_map):
    """Merge all PDF files into a single PDF. OFD files are skipped."""
    merged = fitz.open()
    for fname, fbytes in file_map.items():
        ext = fname.rsplit(".", 1)[-1].lower()
        if ext != "pdf":
            continue
        try:
            doc = fitz.open(stream=fbytes, filetype="pdf")
            merged.insert_pdf(doc)
            doc.close()
        except Exception:
            continue
    if len(merged) == 0:
        return b""
    buf = io.BytesIO()
    merged.save(buf)
    merged.close()
    buf.seek(0)
    return buf.getvalue()


# ─── Smart teacher info parser ──────────────────────
# Auto-detect field types by pattern, regardless of order.

PAT_NAME = re.compile(r'^[一-鿿]{2,4}$')
PAT_ID   = re.compile(r'^\d{17}[\dXx]$')
PAT_CARD = re.compile(r'^\d{15,19}$')
PAT_PHONE = re.compile(r'^1\d{10}$')
PAT_DATE_OR_NUM = re.compile(r'^\d+(\.\d+)?$')

def classify_field(val: str) -> str:
    """Classify a single field value by its format"""
    if PAT_NAME.match(val):
        return "姓名"
    if PAT_ID.match(val):
        return "身份证号"
    if PAT_PHONE.match(val):
        return "手机号"
    if PAT_CARD.match(val):
        return "银行卡号"
    if PAT_DATE_OR_NUM.match(val):
        return "数字"
    return "其他"


def parse_teacher_info(text: str) -> List[Dict]:
    """Parse pasted teacher info, auto-detect fields by pattern regardless of column order"""
    lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
    teachers = []

    for line in lines:
        parts = re.split(r"[\t,，、；;|\s]{1,}", line)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) < 2:
            continue

        teacher = {}
        num_values = []
        for p in parts:
            ftype = classify_field(p)
            if ftype == "姓名":
                if "姓名" not in teacher:
                    teacher["姓名"] = p
            elif ftype == "身份证号":
                teacher["身份证号"] = p
            elif ftype == "手机号":
                teacher["手机号"] = p
            elif ftype == "银行卡号":
                teacher["银行卡号"] = p
            elif ftype == "数字":
                num_values.append(p)

        if "姓名" in teacher:
            teacher.setdefault("身份证号", "")
            teacher.setdefault("手机号", "")
            teacher.setdefault("银行卡号", "")
            teacher.setdefault("标准", 0)
            teacher.setdefault("学时", 0)

            # If there were extra numbers and we have 标准/学时 fields, use them
            if num_values and "标准" in teacher:
                nums = [float(x) for x in num_values]
                if len(nums) >= 1:
                    teacher["标准"] = nums[0]
                if len(nums) >= 2:
                    teacher["学时"] = nums[1]

            teachers.append(teacher)
        else:
            # No clear name found — try first field as name if it looks plausible
            first = parts[0]
            # Check if first non-numeric, non-phone field could be name
            if not PAT_PHONE.match(first) and not PAT_ID.match(first) and not PAT_CARD.match(first):
                teacher["姓名"] = first
                # Re-classify remaining
                for p in parts[1:]:
                    ftype = classify_field(p)
                    if ftype == "身份证号":
                        teacher["身份证号"] = p
                    elif ftype == "手机号":
                        teacher["手机号"] = p
                    elif ftype == "银行卡号":
                        teacher["银行卡号"] = p
                    elif ftype == "数字":
                        if not teacher.get("标准"):
                            teacher["标准"] = float(p)
                        else:
                            teacher["学时"] = float(p)
                teacher.setdefault("身份证号", "")
                teacher.setdefault("手机号", "")
                teacher.setdefault("银行卡号", "")
                teachers.append(teacher)

    return teachers


def parse_teacher_info_basic(text: str) -> List[Dict]:
    """Parse teacher info for travel module (name + bank card, no 标准/学时)"""
    teachers = parse_teacher_info(text)
    # Strip labor-specific fields
    for t in teachers:
        t.pop("标准", None)
        t.pop("学时", None)
    return teachers


# ─── Teacher database from Excel ──────────────────────

def load_teacher_db_from_excel(file_bytes: bytes) -> List[Dict]:
    """Load teacher database from an Excel file. Expected columns: 姓名, 身份证号, 银行卡号, 手机号"""
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    # Read header row
    headers = []
    for cell in ws[1]:
        if cell.value:
            headers.append(str(cell.value).strip())
        else:
            headers.append("")

    # Map columns
    col_map = {}
    for i, h in enumerate(headers):
        for key in ["姓名", "身份证号", "银行卡号", "手机号"]:
            if key in h:
                col_map[key] = i
                break

    teachers = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        teacher = {}
        for key, idx in col_map.items():
            val = row[idx]
            if val is not None and key in ["银行卡号", "身份证号", "手机号"]:
                teacher[key] = str(int(val)) if isinstance(val, (int, float)) else str(val).strip()
            elif val is not None:
                teacher[key] = str(val).strip()
            else:
                teacher[key] = ""
        if "姓名" in teacher and teacher["姓名"]:
            teachers.append(teacher)

    return teachers


# ─── Excel generators ──────────────────────────────

def generate_travel_excel(teachers: List[Dict], project_name: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "差旅费发放表"

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')
    title_font = Font(name='宋体', bold=True, size=14)
    header_font = Font(name='宋体', bold=True, size=11)
    normal_font = Font(name='宋体', size=11)

    max_details = max((len(t.get("明细金额", [])) for t in teachers), default=0)
    total_cols = 4 + max_details

    ws.merge_cells(f'A1:{get_column_letter(total_cols)}1')
    ws['A1'] = "差旅费发放表"
    ws['A1'].font = title_font
    ws['A1'].alignment = center

    ws.merge_cells(f'A2:{get_column_letter(total_cols)}2')
    ws['A2'] = f"项目名称：{project_name}"
    ws['A2'].font = Font(name='宋体', bold=True, size=11)
    ws['A2'].alignment = left_align

    headers = ["序号", "姓名", "银行借记卡号", "实发金额"]
    for col_idx in range(max_details):
        headers.append(f"票面{col_idx+1}")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    total_amount = 0
    for i, t in enumerate(teachers, 1):
        row_num = 3 + i
        ws.row_dimensions[row_num].height = 30
        ws.cell(row=row_num, column=1, value=i).font = normal_font
        ws.cell(row=row_num, column=1).alignment = center
        ws.cell(row=row_num, column=1).border = thin_border

        ws.cell(row=row_num, column=2, value=t.get("姓名", "")).font = normal_font
        ws.cell(row=row_num, column=2).alignment = center
        ws.cell(row=row_num, column=2).border = thin_border

        ws.cell(row=row_num, column=3, value=t.get("银行卡号", "")).font = normal_font
        ws.cell(row=row_num, column=3).alignment = center
        ws.cell(row=row_num, column=3).border = thin_border

        amt = float(t.get("实发金额", 0))
        ws.cell(row=row_num, column=4, value=amt).font = normal_font
        ws.cell(row=row_num, column=4).alignment = center
        ws.cell(row=row_num, column=4).border = thin_border
        ws.cell(row=row_num, column=4).number_format = '#,##0.00'
        total_amount += amt

        # Detail amounts
        details = t.get("明细金额", [])
        for j, d_amt in enumerate(details):
            cell = ws.cell(row=row_num, column=5+j, value=float(d_amt))
            cell.font = normal_font
            cell.alignment = center
            cell.border = thin_border
            cell.number_format = '#,##0.00'

    # Summary
    summary_row = 4 + len(teachers)
    ws.merge_cells(f'A{summary_row}:C{summary_row}')
    ws.cell(row=summary_row, column=1, value="合    计").font = Font(name='宋体', bold=True, size=11)
    ws.cell(row=summary_row, column=1).alignment = center
    ws.cell(row=summary_row, column=1).border = thin_border
    for c in range(2, total_cols + 1):
        ws.cell(row=summary_row, column=c).border = thin_border
    total_cell = ws.cell(row=summary_row, column=4, value=total_amount)
    total_cell.font = Font(name='宋体', bold=True, size=11)
    total_cell.alignment = center
    total_cell.number_format = '#,##0.00'
    total_cell.border = thin_border

    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 24
    ws.column_dimensions['D'].width = 14
    for col_idx in range(max_details):
        ws.column_dimensions[get_column_letter(5+col_idx)].width = 12

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_labor_excel(teachers: List[Dict], project_name: str, project_code: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "劳务费发放表"

    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin'))
    center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    left_align = Alignment(horizontal='left', vertical='center')
    title_font = Font(name='宋体', bold=True, size=14)
    header_font = Font(name='宋体', bold=True, size=10)
    normal_font = Font(name='宋体', size=10)

    ws.merge_cells('A1:J1')
    ws['A1'] = "劳务费发放表"
    ws['A1'].font = title_font
    ws['A1'].alignment = center

    ws.merge_cells('A2:G2')
    ws['A2'] = f"项目名称：{project_name}"
    ws['A2'].font = Font(name='宋体', bold=True, size=10)
    ws['A2'].alignment = left_align
    ws.merge_cells('H2:J2')
    ws['H2'] = f"项目号：{project_code}"
    ws['H2'].font = Font(name='宋体', bold=True, size=10)
    ws['H2'].alignment = left_align

    headers = ["序号", "项目执行时间", "工作内容", "手机号", "姓名", "身份证号", "银行卡号", "标准(元/学时)", "学时", "实发金额"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=h)
        cell.font = header_font
        cell.alignment = center
        cell.border = thin_border

    total_amount = 0
    for i, t in enumerate(teachers, 1):
        row_num = 3 + i
        ws.row_dimensions[row_num].height = 30
        vals = [
            i,
            t.get("执行时间", ""),
            t.get("工作内容", "讲课"),
            t.get("手机号", ""),
            t.get("姓名", ""),
            t.get("身份证号", ""),
            t.get("银行卡号", ""),
            float(t.get("标准", 0)),
            float(t.get("学时", 0)),
        ]
        rate = float(t.get("标准", 0))
        hours = float(t.get("学时", 0))
        amt = rate * hours
        vals.append(amt)

        for col, val in enumerate(vals, 1):
            cell = ws.cell(row=row_num, column=col, value=val)
            cell.font = normal_font
            cell.alignment = center
            cell.border = thin_border
            if col in (8, 10):
                cell.number_format = '#,##0.00'
        total_amount += amt

    summary_row = 4 + len(teachers)
    ws.merge_cells(f'A{summary_row}:G{summary_row}')
    ws.cell(row=summary_row, column=1, value="合    计").font = Font(name='宋体', bold=True, size=10)
    ws.cell(row=summary_row, column=1).alignment = center
    ws.cell(row=summary_row, column=1).border = thin_border
    for c in range(2, 8):
        ws.cell(row=summary_row, column=c).border = thin_border
    total_cell = ws.cell(row=summary_row, column=10, value=total_amount)
    total_cell.font = Font(name='宋体', bold=True, size=10)
    total_cell.alignment = center
    total_cell.number_format = '#,##0.00'
    total_cell.border = thin_border

    sig_row = summary_row + 1
    ws.merge_cells(f'A{sig_row}:J{sig_row}')
    ws.cell(row=sig_row, column=1,
            value="  项目负责人：                           财务负责人：                                      组长：                                  ").font = Font(name='宋体', bold=True, size=10)

    for i, w in enumerate([6, 16, 10, 14, 10, 20, 24, 14, 8, 12], 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── LLM prompts ──────────────────────────────────

TRAVEL_EXTRACT_PROMPT = """你是一个票据处理助手。从上传的文件内容中提取信息。

关键：金额的字段名必须用"票面金额"（即实际金额/价税合计，单位元）。

对每个文件提取：
1. 姓名 - 旅客姓名（优先找旅客姓名，如"刘延川"；保险发票在备注栏找）
2. 日期 - 出发日期（YYYY-MM-DD）
3. 票面金额 - 实际金额（价税合计/合计金额/总金额，数字，单位：元）
4. 文件类型 - "飞机行程单"（有机票/航班信息）或"保险发票"（有保险字样）或"普通发票"

输出纯JSON数组，不要markdown标记：
[
  {
    "姓名": "刘延川",
    "日期": "2026-04-15",
    "票面金额": 880.00,
    "文件类型": "飞机行程单",
    "文件名": "原文件名.pdf"
  }
]"""


# ══════════════════════════════════════════════════
# MAIN APP
# ══════════════════════════════════════════════════

def main():
    st.title("🧾 差旅/劳务费处理系统")

    # Sidebar
    with st.sidebar:
        st.header("配置")
        default_qwen = st.secrets.get("QWEN_API_KEY", "")
        api_key = st.text_input("通义千问 API Key", type="password", value=default_qwen,
                                help="阿里云百炼获取，用于发票识别")
        project_name = st.text_input("项目名称", value="2026-N4-PX")
        project_code = st.text_input("项目号", value="2026-N4-PX")

        # Work log
        st.markdown("---")
        st.markdown("### 工作日志")
        if st.button("清空日志", key="clear_log"):
            st.session_state.work_log = []
            st.rerun()
        log_entries = st.session_state.get("work_log", [])
        if log_entries:
            st.caption(f"共 {len(log_entries)} 条记录")
            for entry in log_entries[-10:]:  # show latest 10
                st.text(f"[{entry.get('time','')}] {entry.get('type','')} {entry.get('summary','')}")
            if len(log_entries) > 10:
                st.caption(f"(更多 {len(log_entries)-10} 条请下载查看)")
            import json
            log_bytes = json.dumps(log_entries, ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button("下载工作日志", data=log_bytes,
                              file_name=f"{project_code}_工作日志.json",
                              mime="application/json")
        else:
            st.caption("暂无记录，生成报表后自动记录")

    # Tab 4 (报道表) doesn't need API key; tabs 1-3 check internally

    # ─── Teacher Database (session-wide) ──────────────
    if "teacher_db" not in st.session_state:
        st.session_state.teacher_db = []

    with st.expander("📚 老师信息库", expanded=not bool(st.session_state.teacher_db)):
        col1, col2 = st.columns([3, 2])
        with col1:
            db_file = st.file_uploader("上传老师信息Excel（姓名/身份证号/银行卡号/手机号）",
                                       type=["xlsx", "xls"], key="db_upload")
        with col2:
            st.markdown("或粘贴更新：")
            db_text = st.text_area("每行一位老师，自动识别字段", height=80,
                                   placeholder="刘延川 6222620910068634421 222325197012140317 13911080938",
                                   key="db_text", label_visibility="collapsed")

        if db_file is not None:
            loaded = load_teacher_db_from_excel(db_file.read())
            if loaded:
                # Merge: update existing, add new
                existing_names = {t["姓名"] for t in st.session_state.teacher_db}
                for t in loaded:
                    if t["姓名"] in existing_names:
                        # Update existing
                        for i, et in enumerate(st.session_state.teacher_db):
                            if et["姓名"] == t["姓名"]:
                                st.session_state.teacher_db[i] = t
                                break
                    else:
                        st.session_state.teacher_db.append(t)
                st.success(f"已更新/合并 {len(loaded)} 位老师")

        if db_text.strip():
            parsed = parse_teacher_info(db_text.strip())
            existing_names = {t["姓名"] for t in st.session_state.teacher_db}
            for t in parsed:
                if t["姓名"] in existing_names:
                    for i, et in enumerate(st.session_state.teacher_db):
                        if et["姓名"] == t["姓名"]:
                            st.session_state.teacher_db[i] = t
                            break
                else:
                    st.session_state.teacher_db.append(t)
            if parsed:
                st.success(f"已更新/新增 {len(parsed)} 位老师")

        if st.session_state.teacher_db:
            st.dataframe(st.session_state.teacher_db, use_container_width=True, hide_index=True)
            st.caption(f"共 {len(st.session_state.teacher_db)} 位老师")
            if st.button("清空信息库", key="clear_db"):
                st.session_state.teacher_db = []
                st.rerun()

    tab1, tab2, tab3, tab4, tab5 = st.tabs(["💰 差旅补助计算", "✈️ 老师差旅报销", "👨‍🏫 老师劳务费发放", "📋 学员报到表", "📜 学员证书发证表"])

    with tab1:
        render_tab_subsidy(api_key, project_name)
    with tab2:
        render_tab_travel(api_key, project_name, project_code)
    with tab3:
        render_tab_labor(project_name, project_code)
    with tab4:
        render_tab_report(project_name)
    with tab5:
        render_tab_certificate()


def render_tab_subsidy(api_key, project_name):
    """原差旅补助计算——名字直接用，不加前缀"""
    st.subheader("差旅补助计算")

    uploaded_files = st.file_uploader("上传PDF行程单", type=["pdf"], accept_multiple_files=True, key="subsidy_files")
    if not uploaded_files:
        st.info("请上传PDF文件")
        return
    st.success(f"已上传 {len(uploaded_files)} 个文件")

    if not st.button("🚀 开始计算", type="primary", key="subsidy_btn"):
        return
    if not api_key:
        st.error("请先在左侧输入通义千问 API Key")
        return

    progress_bar = st.progress(0, text="解析PDF中...")
    status_text = st.empty()

    pdf_texts = []
    file_map = {}
    for i, f in enumerate(uploaded_files):
        bytes_data = f.read()
        file_map[f.name] = bytes_data
        text = smart_extract_pdf(bytes_data, api_key)
        pdf_texts.append(f"【文件{i+1}: {f.name}】\n{text}")
        progress_bar.progress((i + 1) / (len(uploaded_files) + 2))

    combined_text = "\n\n---\n\n".join(pdf_texts)
    status_text.text("调用通义千问...")
    progress_bar.progress(0.6)

    sys_prompt = """从PDF行程单中提取信息，不要计算。
对每个文件提取：
1. 姓名 - 旅客姓名
2. 日期 - 出发日期（YYYY-MM-DD）
3. 票面金额 - 合计金额（数字，元）
4. 类型 - "出发"或"返程"或"中转"
5. 文件名 - 原文件名
输出纯JSON数组。"""
    user_prompt = f"项目名称：{project_name}\n\n文件内容：\n{combined_text}"

    try:
        raw = call_qwen(api_key, sys_prompt, user_prompt)
    except Exception as e:
        st.error(f"API调用失败: {e}")
        return

    progress_bar.progress(0.8)
    try:
        data = parse_json_response(raw)
    except Exception as e:
        st.error(f"解析失败: {e}")
        st.code(raw[:2000])
        return

    # Python calculation
    persons = {}
    for item in data:
        name = item.get("姓名", "未知")
        persons.setdefault(name, []).append(item)

    results = []
    total_fare = total_base = total_extra = total_all = 0
    rename_map = {}

    for person_name, tickets in persons.items():
        tickets.sort(key=lambda x: x.get("日期", ""))
        dates = []
        d_count = r_count = t_count = 0
        for tk in tickets:
            try:
                dates.append(datetime.strptime(str(tk.get("日期", "")), "%Y-%m-%d").date())
            except:
                pass
            tp = tk.get("类型", "")
            if tp == "出发": d_count += 1
            elif tp == "返程": r_count += 1
            else: t_count += 1
        if not dates:
            continue
        travel_days = (max(dates) - min(dates)).days + 1
        base = travel_days * 100
        extra = (d_count + r_count + t_count) * 80
        total_person = base + extra

        for tk in tickets:
            fare = float(tk.get("票面金额", 0))
            d_str = str(tk.get("日期", ""))
            try:
                dt = datetime.strptime(d_str, "%Y-%m-%d")
                disp_date = f"{dt.month}月{dt.day}日"
            except:
                disp_date = d_str
            new_name = f"{person_name} {disp_date} {fare:.2f}.pdf"
            results.append({
                "姓名": person_name, "日期": d_str, "票面金额": fare,
                "基础补助": base, "额外补助": extra, "总金额": total_person,
                "原文件名": tk.get("文件名", ""), "新文件名": new_name
            })
            orig = tk.get("文件名", "")
            if orig and orig in file_map:
                rename_map[orig] = new_name
            total_fare += fare
        total_base += base
        total_extra += extra
        total_all += total_person

    progress_bar.progress(1.0)
    status_text.text("✅ 完成")

    log_work("差旅补助", f"{len(results)} 人, 总¥{total_all:.2f}", {"人数": len(persons), "金额": total_all})
    st.subheader("明细")
    st.dataframe([{
        "姓名": r["姓名"], "日期": r["日期"], "票面金额": r["票面金额"],
        "基础补助": r["基础补助"], "额外补助": r["额外补助"],
        "总金额": r["总金额"], "新文件名": r["新文件名"]
    } for r in results], use_container_width=True, hide_index=True)

    st.subheader("💰 总计")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("票面金额合计", f"¥{total_fare:.2f}")
    c2.metric("基础补助合计", f"¥{total_base:.2f}")
    c3.metric("额外补助合计", f"¥{total_extra:.2f}")
    c4.metric("总金额合计", f"¥{total_all:.2f}")

    if rename_map:
        st.subheader("📁 文件改名")
        st.dataframe([{"原文件名": k, "新文件名": v} for k, v in rename_map.items()],
                     use_container_width=True, hide_index=True)
        zip_bytes = create_renamed_zip(file_map, rename_map)
        st.download_button("📦 下载改名文件(ZIP)", data=zip_bytes,
                          file_name=f"{project_name}_改名文件.zip", mime="application/zip")


def render_tab_travel(api_key, project_name, project_code):
    """老师差旅报销——Excel姓名写老师原名，文件改名加'段'前缀"""
    st.subheader("老师差旅报销")
    st.markdown("""
    - 实发金额 = 票面金额合计，含保险费则相加
    - 普通发票从文件名提取姓名
    - **改名格式：`段 {姓名} {日期} {金额}.pdf`**
    """)

    uploaded_files = st.file_uploader("上传PDF/OFD文件", type=["pdf", "ofd"],
                                      accept_multiple_files=True, key="travel_files")
    if not uploaded_files:
        st.info("请上传文件")
        return
    st.success(f"已上传 {len(uploaded_files)} 个文件")

    # Bank info — use teacher DB or paste
    st.markdown("---")
    st.markdown("**老师银行卡信息**")
    use_db = st.checkbox("从老师信息库自动匹配", value=bool(st.session_state.teacher_db))

    teachers = []
    if use_db and st.session_state.teacher_db:
        teachers = st.session_state.teacher_db.copy()
        st.success(f"已加载信息库 {len(teachers)} 位老师")
    else:
        bank_text = st.text_area("粘贴（姓名+银行卡号，每行一位）",
                                 placeholder="刘延川 6222620910068634421",
                                 height=100, key="travel_bank")
        if bank_text.strip():
            teachers = parse_teacher_info(bank_text.strip())

    if not st.button("🚀 处理并生成", type="primary", key="travel_btn"):
        return
    if not api_key:
        st.error("请先在左侧输入通义千问 API Key")
        return
    if not teachers:
        st.info("未提供老师信息，将直接从文件提取（银行卡号为空）")

    progress_bar = st.progress(0, text="解析文件中...")
    status_text = st.empty()

    file_data = []
    file_map = {}
    pdf_texts = []

    for i, f in enumerate(uploaded_files):
        bytes_data = f.read()
        file_map[f.name] = bytes_data
        ext = Path(f.name).suffix.lower()

        if ext == ".ofd":
            parsed = parse_ofd(bytes_data)
            file_data.append({
                "文件名": f.name, "姓名": parsed.get("姓名", ""),
                "金额": parsed.get("金额", 0), "保险费": parsed.get("保险费", 0),
                "日期": parsed.get("日期", ""), "类型": "ofd",
                "原始文本": parsed.get("原始文本", "")
            })
        else:
            text = smart_extract_pdf(bytes_data, api_key)
            pdf_texts.append(f"【文件{i+1}: {f.name}】\n{text}")
            file_data.append({"文件名": f.name, "类型": "pdf", "原始文本": text})
        progress_bar.progress((i + 1) / (len(uploaded_files) + 2))

    # LLM for PDFs
    if pdf_texts:
        status_text.text("调用通义千问解析PDF...")
        progress_bar.progress(0.6)
        combined = "\n\n---\n\n".join(pdf_texts)
        try:
            llm_raw = call_qwen(api_key, TRAVEL_EXTRACT_PROMPT,
                                f"项目名称：{project_name}\n\n文件内容：\n{combined}")
            llm_data = parse_json_response(llm_raw)
            for item in llm_data:
                for fd in file_data:
                    if fd["文件名"] == item.get("文件名", "") and fd["类型"] == "pdf":
                        fd["姓名"] = item.get("姓名", "")
                        fd["金额"] = float(item.get("票面金额", 0))
                        fd["保险费"] = float(item.get("保险费", 0))
                        fd["日期"] = item.get("日期", "")
                        fd["文件类型"] = item.get("文件类型", "普通发票")
                        break
        except Exception as e:
            st.warning(f"LLM解析部分失败: {e}，从文件名提取")

    # Fallback: extract names from filenames
    PLACE_WORDS = {"大连","郑州","北京","上海","广州","深圳","天津","重庆",
                   "南京","西安","杭州","成都","武汉","昆明","厦门","长沙"}
    SKIP_WORDS = PLACE_WORDS | {"保险发票","行程单","电子发票","航空运输","客票行程单"}
    for fd in file_data:
        if not fd.get("姓名"):
            stem = Path(fd["文件名"]).stem
            m = re.search(r'_([一-鿿]{2,4})\s+\d{2}月', stem)
            if m:
                fd["姓名"] = m.group(1)
            else:
                cn_names = re.findall(r'[一-鿿]{2,4}', stem)
                clean = [n for n in cn_names if n not in SKIP_WORDS]
                fd["姓名"] = clean[-1] if clean else "未知"

    # Correct LLM-place names with filename
    teacher_names = {t["姓名"] for t in teachers}
    for fd in file_data:
        llm_name = fd.get("姓名", "")
        if llm_name and llm_name not in teacher_names:
            stem = Path(fd["文件名"]).stem
            m = re.search(r'_([一-鿿]{2,4})\s+\d{2}月', stem)
            if m and m.group(1) in teacher_names:
                fd["姓名"] = m.group(1)

    progress_bar.progress(0.8, text="匹配老师和金额...")

    # Match
    teacher_amounts = {}
    teacher_details = {}
    teacher_dates = {}
    for t in teachers:
        n = t["姓名"]
        teacher_amounts[n] = 0.0
        teacher_details[n] = []
        teacher_dates[n] = []

    for fd in file_data:
        p_name = fd.get("姓名", "")
        amt = fd.get("金额", 0)
        date_str = fd.get("日期", "")

        matched = False
        for t in teachers:
            t_name = t["姓名"]
            if p_name and (p_name in t_name or t_name in p_name):
                teacher_amounts[t_name] += amt
                teacher_details[t_name].append(amt)
                if date_str:
                    teacher_dates[t_name].append(date_str)
                matched = True
                break
        if not matched and p_name and amt > 0:
            teachers.append({"姓名": p_name, "银行卡号": ""})
            teacher_amounts[p_name] = amt
            teacher_details[p_name] = [amt]
            if date_str:
                teacher_dates[p_name] = [date_str]

    progress_bar.progress(0.9, text="生成结果...")

    # Build output: Excel uses teacher name only; rename files use "段 {name}"
    output_teachers = []
    rename_map = {}
    file_lookup = {fd["文件名"]: fd for fd in file_data}
    total = 0

    for t in teachers:
        name = t["姓名"]
        amt = teacher_amounts.get(name, 0)
        detail = teacher_details.get(name, [])
        output_teachers.append({
            "姓名": name,
            "银行卡号": t.get("银行卡号", ""),
            "实发金额": amt,
            "明细金额": detail
        })
        total += amt

    # Build rename map: for each uploaded file, find matching teacher and create new name
    for orig_name in file_map:
        fd = file_lookup.get(orig_name)
        if fd:
            p_name = fd.get("姓名", "")
            # Find matching teacher
            teacher_name = None
            for t in teachers:
                if p_name and (p_name in t["姓名"] or t["姓名"] in p_name):
                    teacher_name = t["姓名"]
                    break
            if not teacher_name:
                teacher_name = p_name

            amt = fd.get("金额", 0)
            date_str = fd.get("日期", "")
            if date_str:
                try:
                    dt = datetime.strptime(date_str, "%Y-%m-%d")
                    disp_date = f"{dt.month}月{dt.day}日"
                except:
                    disp_date = date_str
            else:
                # Try to get from filename
                disp_date = ""

            new_name = f"段 {teacher_name} {disp_date} {amt:.2f}.pdf"
            rename_map[orig_name] = new_name

    progress_bar.progress(1.0)
    status_text.text("✅ 完成")

    # Display results
    log_work("老师差旅报销", f"{len(output_teachers)} 人, 总¥{total:.2f}", {"人数": len(output_teachers), "金额": total})
    st.subheader("📊 处理结果")
    table_rows = []
    for ot in output_teachers:
        detail_str = " + ".join([f"{v:.2f}" for v in ot["明细金额"]])
        table_rows.append({
            "姓名": ot["姓名"], "银行卡号": ot["银行卡号"],
            "实发金额": ot["实发金额"], "明细": detail_str
        })
    st.dataframe(table_rows, use_container_width=True, hide_index=True)
    st.metric("实发总金额", f"¥{total:.2f}")

    # Excel download
    col1, col2 = st.columns(2)
    with col1:
        excel_bytes = generate_travel_excel(output_teachers, project_name)
        st.download_button("📥 下载差旅费Excel", data=excel_bytes,
                          file_name=f"{project_name}_差旅费.xlsx",
                          mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                          type="primary")

    # ZIP download for renamed files
    with col2:
        if rename_map:
            st.dataframe([{"原文件名": k, "新文件名": v} for k, v in rename_map.items()],
                         use_container_width=True, hide_index=True)
            zip_bytes = create_renamed_zip(file_map, rename_map)
            st.download_button("📦 下载改名文件(ZIP)", data=zip_bytes,
                              file_name=f"{project_name}_改名文件.zip",
                              mime="application/zip", type="primary")

    # PDF merge
    st.markdown("---")
    st.subheader("📄 PDF 合集（仅合并 PDF 文件，OFD 跳过）")
    merged_pdf = merge_pdfs(file_map)
    if merged_pdf:
        st.download_button("📥 下载 PDF 合集", data=merged_pdf,
                          file_name=f"{project_name}_发票合集.pdf",
                          mime="application/pdf", type="primary")
    else:
        st.info("上传文件中没有 PDF，无法生成合集")


def render_tab_labor(project_name, project_code):
    """老师劳务费发放——从信息库拉取，或粘贴完整信息"""
    st.subheader("老师劳务费发放")
    st.markdown("从信息库拉取或粘贴老师信息")

    # Use teacher DB or paste
    use_db = st.checkbox("从老师信息库拉取银行卡号/身份证号/手机号",
                         value=bool(st.session_state.teacher_db), key="labor_use_db")

    # Default values for 标准 and 学时
    col_a, col_b = st.columns(2)
    with col_a:
        default_rate = st.number_input("默认标准（元/学时）", min_value=0.0, value=1000.0, step=100.0, key="labor_rate")
    with col_b:
        default_hours = st.number_input("默认学时", min_value=0.0, value=8.0, step=1.0, key="labor_hours")

    info_text = st.text_area(
        "粘贴数据（只需姓名和执行时间，标准和学时用默认值，有则覆盖）",
        placeholder="李阳  5月13日上午  讲课  13911080938  222325197012140317  6222600910058226273\n王芳  5月15日全天  讲课  13601023762  110108196512152251  9558800200125094454\n（或每行末尾加数字：姓名 时间 标准 学时 → 覆盖默认值）",
        height=150, key="labor_info"
    )

    teachers = []
    if info_text.strip():
        teachers = parse_teacher_info(info_text.strip())
        # Apply defaults if not set by paste
        for t in teachers:
            if not t.get("标准") or t["标准"] == 0:
                t["标准"] = default_rate
            if not t.get("学时") or t["学时"] == 0:
                t["学时"] = default_hours

        # Merge with DB
        if use_db and st.session_state.teacher_db:
            db_lookup = {t["姓名"]: t for t in st.session_state.teacher_db}
            for t in teachers:
                db_entry = db_lookup.get(t["姓名"])
                if db_entry:
                    if not t.get("手机号") and db_entry.get("手机号"):
                        t["手机号"] = db_entry["手机号"]
                    if not t.get("身份证号") and db_entry.get("身份证号"):
                        t["身份证号"] = db_entry["身份证号"]
                    if not t.get("银行卡号") and db_entry.get("银行卡号"):
                        t["银行卡号"] = db_entry["银行卡号"]

        if teachers:
            st.dataframe(teachers, use_container_width=True, hide_index=True)
            st.success(f"识别到 {len(teachers)} 位老师")
        else:
            st.warning("未识别到有效数据")
            teachers = []

    if st.button("🚀 生成劳务费Excel", type="primary", key="labor_btn"):
        if not teachers:
            st.warning("请先粘贴有效信息")
            return

        excel_bytes = generate_labor_excel(teachers, project_name, project_code)
        total = sum(float(t.get("标准", 0)) * float(t.get("学时", 0)) for t in teachers)
        log_work("劳务费发放", f"{len(teachers)} 人, 总¥{total:.2f}", {"人数": len(teachers), "金额": total})
        st.metric("总金额", f"¥{total:.2f}")
        st.download_button("📥 下载劳务费Excel", data=excel_bytes,
                          file_name=f"{project_name}_劳务费.xlsx",
                          mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                          type="primary")


def generate_certificate_excel(students, cert_date, params):
    """Generate 2-sheet certificate Excel matching the template format."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    ws1 = wb.active
    ws1.title = "发证"
    ws1.page_setup.paperSize = 9
    ws1.page_setup.orientation = "portrait"
    ws1.page_margins.left = 0.4
    ws1.page_margins.right = 0.4
    ws1.page_margins.top = 0.5
    ws1.page_margins.bottom = 0.5
    ws1.print_title_rows = "1:1"

    thin = Border(left=Side("thin"), right=Side("thin"),
                  top=Side("thin"), bottom=Side("thin"))
    center = Alignment(horizontal="center", vertical="center")
    wrap = Alignment(horizontal="center", vertical="center", wrap_text=True)
    hdr_font = Font(name="宋体", bold=True, size=11)
    norm_font = Font(name="宋体", size=11)

    h1 = ["序号", "姓名", "身份证", "证书名称", "培训日期", "项目名称", "培训地点",
          "起始证书有效期", "证书截止有效期", "证书编号", "学时", "成绩",
          "是否公示", "发证单位", "发证日期", "上传部门", "负责人"]
    ws1_ws = [6.63, 8.63, 19.45, 10.82, 23.27, 52.73, 9.82,
              17.27, 15.27, 22.63, 9.27, 13.0, 13.0, 21.54, 13.82, 15.18, 16.45]

    for ci, h in enumerate(h1, 1):
        cell = ws1.cell(row=1, column=ci, value=h)
        cell.font = hdr_font
        cell.alignment = wrap if h in ("身份证",) else center
        cell.border = thin
        ws1.column_dimensions[get_column_letter(ci)].width = ws1_ws[ci - 1]
    ws1.row_dimensions[1].height = 24

    year = params["cert_year"]
    seq_start = int(params["cert_seq_start"])
    for i, s in enumerate(students):
        row = 2 + i
        ws1.row_dimensions[row].height = 24
        cert_no = "CASEI-HYPX-" + year + "-" + str(seq_start + i).zfill(4)
        vals = [
            i + 1,
            s.get("姓名", ""),
            s.get("身份证", ""),
            "培训证明",
            params.get("train_date", ""),
            params.get("project_name", ""),
            params.get("train_location", ""),
            "",
            "长期",
            cert_no,
            params.get("hours", ""),
            "",
            "否",
            "中国特种设备检验协会",
            cert_date,
            "",
            params.get("person_in_charge", "段秉泽"),
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws1.cell(row=row, column=ci, value=v)
            cell.font = norm_font
            cell.alignment = wrap if h1[ci - 1] in ("身份证",) else center
            cell.border = thin
            if ci == 15:
                cell.number_format = '@'

    ws2 = wb.create_sheet("打印")
    ws2.page_setup.paperSize = 9
    ws2.page_setup.orientation = "landscape"
    ws2.page_margins.left = 0.4
    ws2.page_margins.right = 0.4
    ws2.page_margins.top = 0.5
    ws2.page_margins.bottom = 0.5
    ws2.print_title_rows = "1:1"

    hdr_sizes = [10, 12, 10, 11, 10, 10, 11, 10]
    h2 = ["序号", "省份", "姓名", "身份证", "手机", "单位", "证书编号", "备注"]
    ws2_ws = [6.63, 10.0, 8.63, 19.45, 15.0, 36.27, 22.63, 8.0]

    for ci, h in enumerate(h2, 1):
        cell = ws2.cell(row=1, column=ci, value=h)
        cell.font = Font(name="宋体", bold=True, size=hdr_sizes[ci - 1])
        cell.alignment = wrap
        cell.border = thin
        ws2.column_dimensions[get_column_letter(ci)].width = ws2_ws[ci - 1]
    ws2.row_dimensions[1].height = 30

    for i, s in enumerate(students):
        row = 2 + i
        ws2.row_dimensions[row].height = 30
        cert_no = "CASEI-HYPX-" + year + "-" + str(seq_start + i).zfill(4)
        vals = [
            i + 1,
            s.get("省份", ""),
            s.get("姓名", ""),
            s.get("身份证", ""),
            s.get("手机", ""),
            s.get("纳税人名称", ""),
            cert_no,
            "",
        ]
        for ci, v in enumerate(vals, 1):
            cell = ws2.cell(row=row, column=ci, value=v)
            cell.font = Font(name="宋体", size=10)
            cell.alignment = wrap if ci in (4, 6) else Alignment(horizontal="center", vertical="center")
            cell.border = thin

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─── Word template functions ───────────────────────


def fill_docx_template(template_path: Path, replacements: Dict[str, str]) -> bytes:
    """Read a .docx template, replace placeholder text in paragraphs and tables, return bytes."""
    import docx
    doc = docx.Document(str(template_path))

    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        replaced = full
        for old, new in replacements.items():
            replaced = replaced.replace(old, new)
        if replaced != full and p.runs:
            p.runs[0].text = replaced
            for r in p.runs[1:]:
                r.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full = "".join(r.text for r in p.runs)
                    replaced = full
                    for old, new in replacements.items():
                        replaced = replaced.replace(old, new)
                    if replaced != full and p.runs:
                        p.runs[0].text = replaced
                        for r in p.runs[1:]:
                            r.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_certificate_word(params: dict) -> bytes:
    """Fill the training certificate template with project-level params."""
    from datetime import date
    train_end = params.get("train_date_end", date.today())
    train_start = params.get("train_date_start", date.today())

    # Template format: "2026年04月15日-04月17日" (year only on first date)
    s, e = train_start, train_end
    old_date = f"{s.year}年{s.month:02d}月{s.day:02d}日-{e.month:02d}月{e.day:02d}日"
    new_date = f"{s.year}年{s.month}月{s.day}日-{e.month}月{e.day}日"
    sign_date = f"{train_end.year}年{train_end.month:02d}月"

    template = Path(__file__).parent / "培训证明模板.docx"
    import docx
    doc = docx.Document(str(template))

    # Pass 1: fill training text paragraph (has «姓名» or "参加")
    text_para = None
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if "«姓名»" in full or "参加" in full:
            text_para = p
            replaced = full
            replaced = replaced.replace(old_date, new_date)
            replaced = replaced.replace("第四届起重机械检验检测与安全管理技术交流会", params.get("project_name", ""))
            replaced = replaced.replace("20学时", f"{params.get('hours', '32')}学时")
            if replaced != full and p.runs:
                p.runs[0].text = replaced
                for r in p.runs[1:]:
                    r.text = ""
            break

    # Pass 2: sign-off date paragraph (short, only year+month)
    for p in doc.paragraphs:
        if p is text_para:
            continue
        full = "".join(r.text for r in p.runs)
        txt = full.strip()
        if txt and "年" in txt and txt.count("月") >= 1 and len(txt) <= 10 and "中国" not in txt:
            if p.runs:
                p.runs[0].text = sign_date
                for r in p.runs[1:]:
                    r.text = ""
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            break

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _set_cell_text(cell, text):
    """Clear ALL paragraphs in a cell, set first paragraph to text."""
    for i, p in enumerate(cell.paragraphs):
        for r in p.runs:
            r.text = ""
        if i == 0:
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)


def generate_approval_word(params: dict, student_count: int) -> bytes:
    """Fill the approval form template with table-based replacement."""
    from datetime import date
    train_start = params.get("train_date_start", date.today())
    train_end = params.get("train_date_end", date.today())

    start_str = f"{train_start.year}.{train_start.month:02d}.{train_start.day:02d}"
    end_str = f"{train_end.year}.{train_end.month:02d}.{train_end.day:02d}"
    date_range = f"{start_str}-{end_str}"

    cert_year = params.get("cert_year", "2026")
    cert_start = int(params.get("cert_seq_start", "0001"))
    cert_end = cert_start + student_count - 1
    cert_range = f"CASEI-HYPX-{cert_year}-{str(cert_start).zfill(4)}  --  CASEI-HYPX-{cert_year}-{str(cert_end).zfill(4)}"

    template = Path(__file__).parent / "审批表模板.docx"
    import docx
    doc = docx.Document(str(template))
    table = doc.tables[0]

    # Row 0: 项目名称
    _set_cell_text(table.rows[0].cells[1], params.get("project_name", ""))
    # Row 1: 培训时间 (cell 1), 培训地点 (cell 4)
    _set_cell_text(table.rows[1].cells[1], date_range)
    _set_cell_text(table.rows[1].cells[4], params.get("train_location", ""))
    # Row 2: 项目负责人 (cell 1), 发证数量 (cell 4)
    _set_cell_text(table.rows[2].cells[1], params.get("person_in_charge", ""))
    _set_cell_text(table.rows[2].cells[4], str(student_count))
    # Row 6: 证书号码号段 — merged, write to cell 0
    _set_cell_text(table.rows[6].cells[0], f"证书号码（号段）\n{cert_range}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
def render_tab_certificate():
    """Tab 5: Certificate issuance — unified params, Excel + 2 Word docs"""
    st.subheader("学员证书发证表生成")

    st.markdown("### 1. 上传学员数据")
    src = st.file_uploader("上传报到表 Excel", type=["xlsx", "xls"], key="cert_src")
    students = []
    if src:
        try:
            import pandas as pd
            df = pd.read_excel(src)
            col_map = {}
            for c in df.columns:
                cs = str(c)
                for k in ["省份", "姓名", "身份证", "手机", "纳税人名称"]:
                    if k in cs:
                        col_map[k] = c
            for _, row in df.iterrows():
                s = {}
                for k in col_map:
                    v = row[col_map[k]]
                    if k in ("手机", "身份证") and pd.notna(v) and isinstance(v, float):
                        s[k] = str(int(v))
                    elif pd.notna(v):
                        s[k] = str(v)
                    else:
                        s[k] = ""
                if s.get("姓名"):
                    students.append(s)
            st.success(f"读取到 {len(students)} 名学员")
            st.dataframe(pd.DataFrame(students), use_container_width=True, hide_index=True)
        except Exception as e:
            st.error(f"读取失败: {e}")

    st.markdown("### 2. 填写参数（Excel & Word 共用）")
    col1, col2 = st.columns(2)
    with col1:
        proj = st.text_input("项目名称（完整）",
                             value="第四届起重机械检验检测与安全管理技术交流会",
                             key="cert_proj")
        loc = st.text_input("培训地点", value="开封", key="cert_loc")
    with col2:
        hours = st.text_input("学时", value="20", key="cert_hours")
        person = st.text_input("负责人", value="段秉泽", key="cert_person")

    st.markdown("**培训日期（起止日期）**")
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    with c1:
        s_mon = st.selectbox("起始月", list(range(1, 13)), index=3, key="cert_smon")  # default Apr
    with c2:
        s_day = st.selectbox("起始日", list(range(1, 32)), index=14, key="cert_sday")  # default 15
    with c3:
        e_mon = st.selectbox("结束月", list(range(1, 13)), index=3, key="cert_emon")  # default Apr
    with c4:
        e_day = st.selectbox("结束日", list(range(1, 32)), index=16, key="cert_eday")  # default 17

    st.markdown("**证书编号**")
    cc1, cc2, cc3, cc4 = st.columns(4)
    with cc1:
        cert_year = st.text_input("年份", value="2026", key="cert_year")
    with cc2:
        cert_seq = st.text_input("起始号", placeholder="0530", key="cert_seq")

    st.markdown("**发证日期（Excel 用）**")
    c3, c4 = st.columns(2)
    with c3:
        i_mon = st.selectbox("月", list(range(1, 13)), index=e_mon-1, key="cert_mon")
    with c4:
        i_day = st.selectbox("日", list(range(1, 32)), key="cert_day")

    if st.button("🚀 生成全部", type="primary", key="cert_btn"):
        if not students:
            st.warning("请先上传报到表")
        elif not cert_seq:
            st.warning("请填写证书编号起始号")
        else:
            from datetime import date
            train_start = date(int(cert_year), s_mon, s_day)
            train_end = date(int(cert_year), e_mon, e_day)
            cert_date_str = f"{int(cert_year)}年{i_mon:02d}月{i_day:02d}日"

            params = {
                "project_name": proj, "train_location": loc, "hours": hours,
                "cert_year": cert_year, "cert_seq_start": cert_seq, "person_in_charge": person,
                "train_date_start": train_start, "train_date_end": train_end,
                "train_date": f"{train_start.year}年{train_start.month}月{train_start.day}日-{train_end.month}月{train_end.day}日",
                "cert_date_str": cert_date_str,
            }
            end_seq = int(cert_seq) + len(students) - 1
            safe_name = proj.replace("/", "-").replace("\\", "-")[:40] if proj else "发证表"

            st.session_state.cert_xlsx = generate_certificate_excel(students, cert_date_str, params)
            st.session_state.cert_word = generate_certificate_word(params)
            st.session_state.cert_approval = generate_approval_word(params, len(students))
            st.session_state.cert_safe_name = safe_name
            st.session_state.cert_label = f"CASEI-HYPX-{cert_year}-{cert_seq} ~ CASEI-HYPX-{cert_year}-{str(end_seq).zfill(4)}"
            st.session_state.cert_count = len(students)
            st.session_state.cert_log = f"{len(students)} 人, 编号 {cert_year}-{cert_seq}~{str(end_seq).zfill(4)}"
            st.rerun()

    # Show download buttons from session_state (persist across downloads)
    if st.session_state.get("cert_xlsx"):
        log_work("证书发证表", st.session_state.cert_log)
        st.success(f"生成完成: {st.session_state.cert_count} 条记录")
        st.metric("证书编号范围", st.session_state.cert_label)
        safe_name = st.session_state.cert_safe_name

        dl1, dl2, dl3 = st.columns(3)
        with dl1:
            st.download_button("📥 发证表 Excel", data=st.session_state.cert_xlsx,
                              file_name=f"{safe_name}_证书发证表.xlsx",
                              mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        with dl2:
            st.download_button("📥 培训证明 Word", data=st.session_state.cert_word,
                              file_name=f"{safe_name}_培训证明.docx",
                              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        with dl3:
            st.download_button("📥 审批表 Word", data=st.session_state.cert_approval,
                              file_name=f"{safe_name}_审批表.docx",
                              mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ─── Tab 4: 学员报到表 ─────────────────────────────


def render_tab_report(project_name):
    """学员报到表生成——A3横版打印"""
    st.subheader("学员报到表生成")
    st.markdown("上传学员信息表，选保留列，生成A3横版打印报到表")

    col1, col2 = st.columns(2)
    with col1:
        source_file = st.file_uploader("上传学员信息表", type=["xls", "xlsx"], key="rpt_source")
    with col2:
        member_file = st.file_uploader("会员单位表（可选，内置默认，上传可覆盖）", type=["xls", "xlsx"], key="rpt_member")

    if source_file is None:
        st.info("请上传学员信息表")
        return

    import pandas as pd
    from pathlib import Path

    # Read source
    try:
        df = pd.read_excel(source_file)
    except Exception as e:
        st.error(f"读取文件失败: {e}")
        return

    # Read member list — auto-load from repo, override if user uploads
    member_set = set()
    if "rpt_member_loaded" not in st.session_state:
        bundled = Path(__file__).parent / "会员单位表.xls"
        if bundled.exists():
            try:
                mdf = pd.read_excel(str(bundled))
                for col in mdf.columns:
                    if '单位' in str(col):
                        member_set = set(mdf[col].dropna().astype(str).str.strip().str.replace('　', '').str.replace(' ', ''))
                        break
                st.session_state.rpt_member_loaded = member_set
                st.info(f"已加载内置会员单位 {len(member_set)} 家")
            except:
                st.session_state.rpt_member_loaded = set()

    if member_file:
        try:
            mdf = pd.read_excel(member_file)
            for col in mdf.columns:
                if '单位' in str(col):
                    member_set = set(mdf[col].dropna().astype(str).str.strip().str.replace('　', '').str.replace(' ', ''))
                    break
            st.success(f"已覆盖加载会员单位 {len(member_set)} 家")
        except Exception as e:
            st.warning(f"会员表读取失败（不影响主功能）: {e}")

    # Fallback to session_state (auto-loaded) if no upload
    if not member_set and st.session_state.get("rpt_member_loaded"):
        member_set = st.session_state.rpt_member_loaded

    # Column picker
    all_cols = list(df.columns)
    st.markdown("### 选择保留的列")

    # Suggest defaults: match user's desired order
    desired_order = ['序号','省份','姓名','身份证','手机','发票类型','纳税人名称','纳税人识别号','邮箱']
    suggest = [c for c in desired_order if c in all_cols]
    selected = st.multiselect("勾选需要显示的列", all_cols, default=suggest[:13])

    if not selected:
        st.warning("至少选一列")
        return

    # Custom title
    title = st.text_input("页眉大标题", value="学员报到表")

    # Generate
    if not st.button("📄 生成报到表", type="primary", key="rpt_btn"):
        return

    # Find unit column for member check (单位名称 or 纳税人名称)
    unit_col = None
    for c in selected:
        if '单位' in str(c) or '纳税' in str(c):
            unit_col = c
            break

    # Build rows
    display_df = df[selected].copy()

    # Dedup by name + ID
    if '姓名' in display_df.columns and '身份证' in display_df.columns:
        display_df = display_df.drop_duplicates(subset=['姓名', '身份证'], keep='first').reset_index(drop=True)

    # Sort: by 省份 first, then by 纳税人名称/单位名称
    sort_cols = []
    for c in selected:
        if '省份' in str(c) or '地区' in str(c) or str(c).strip() == '省':
            sort_cols.append(c)
            break
    for c in selected:
        if '纳税' in str(c) or '单位' in str(c):
            sort_cols.append(c)
            break
    if sort_cols:
        display_df = display_df.sort_values(by=sort_cols, na_position='last').reset_index(drop=True)

    # Auto-fill 序号 column
    for c in selected:
        if '序号' in str(c):
            display_df[c] = range(1, len(display_df) + 1)
            break

    member_rows = set()
    if unit_col and member_set:
        for idx, row in display_df.iterrows():
            unit_val = str(row[unit_col]).strip().replace('　', '').replace(' ', '')
            if unit_val in member_set:
                member_rows.add(idx)

    n_data = len(display_df)
    n_blank = 50
    total = n_data + n_blank

    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "报到表"

    # Page setup: A3 landscape
    ws.page_setup.paperSize = 8
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins.left = 1.0
    ws.page_margins.right = 1.0
    ws.page_margins.top = 1.0
    ws.page_margins.bottom = 1.0

    # Repeat header row on each printed page
    ws.print_title_rows = '1:1'

    thin = Border(left=Side('thin'), right=Side('thin'),
                  top=Side('thin'), bottom=Side('thin'))
    center_wrap = Alignment(horizontal='center', vertical='center', wrap_text=True)
    title_font = Font(name='宋体', bold=True, size=16)
    header_font = Font(name='宋体', bold=True, size=10)
    normal_font = Font(name='宋体', size=10)
    bold_font = Font(name='宋体', bold=True, size=10)

    # Add blank columns for manual entry
    extra_cols = ["本人签字", "付款方式", "金额", "备注"]
    all_headers = list(selected) + extra_cols
    num_cols = len(all_headers)

    # Column widths: auto for selected, fixed for extras
    col_widths = {}
    for i, h in enumerate(selected):
        name = str(h)
        if '序号' in name: col_widths[i] = 6
        elif '省份' in name or '地区' in name: col_widths[i] = 8
        elif '姓名' in name: col_widths[i] = 10
        elif '身份证' in name: col_widths[i] = 22
        elif '手机' in name or '电话' in name: col_widths[i] = 14
        elif '单位' in name: col_widths[i] = 22
        elif '发票' in name: col_widths[i] = 8
        elif '纳税' in name: col_widths[i] = 22
        elif '邮箱' in name: col_widths[i] = 20
        else: col_widths[i] = 12

    # Widths for extra columns
    extra_widths = [10, 10, 10, 14]

    # Print header (doesn't block Excel filtering)
    ws.oddHeader.center.text = f"&20&B{title}"

    # Row 1: Headers
    for ci, h in enumerate(all_headers, 1):
        cell = ws.cell(row=1, column=ci, value=str(h))
        cell.font = header_font
        cell.alignment = center_wrap
        cell.border = thin
        # Set column width
        if ci <= len(selected):
            ws.column_dimensions[get_column_letter(ci)].width = col_widths.get(ci-1, 12)
        else:
            ei = ci - len(selected) - 1
            ws.column_dimensions[get_column_letter(ci)].width = extra_widths[ei] if ei < len(extra_widths) else 12
    ws.row_dimensions[1].height = 30

    # Data + blank rows
    for ri in range(total):
        row_num = 2 + ri
        ws.row_dimensions[row_num].height = 35
        if ri < n_data:
            row = display_df.iloc[ri]
            is_member = ri in member_rows
            fnt = bold_font if is_member else normal_font
            for ci in range(len(selected)):
                val = row.iloc[ci]
                display = str(val) if pd.notna(val) else ""
                cell = ws.cell(row=row_num, column=ci+1, value=display)
                cell.font = fnt
                cell.alignment = center_wrap
                cell.border = thin
            # Blank extra columns
            for ci in range(len(extra_cols)):
                cell = ws.cell(row=row_num, column=len(selected)+ci+1, value="")
                cell.font = fnt
                cell.alignment = center_wrap
                cell.border = thin
        else:
            for ci in range(num_cols):
                cell = ws.cell(row=row_num, column=ci+1, value="")
                cell.alignment = center_wrap
                cell.border = thin

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    log_work("学员报到表", f"{n_data} 人")
    st.success(f"生成完成：实际 {n_data} 人，总行 {total} 行（含 {n_blank} 行空白）")
    st.download_button("📥 下载报到表 Excel", data=buf,
                       file_name=f"{project_name}_{title}.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                       type="primary")


if __name__ == "__main__":
    main()
